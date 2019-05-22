import sqlite3
import os
import docx
import re

conn = sqlite3.connect('entries.db')
cursor = conn.cursor()

# reset database for testing purposes
cursor.execute("DELETE FROM entries")

# regex for matching my date formats
def date_match(strg, search=re.compile(r'[^0-9./\- ]').search):
    return strg != '' and strg != '---' and not bool(search(strg))

# parse folder of .docx files
for file in os.listdir('./entry_files'):
    document = docx.Document('./entry_files/'+file)
    current_date = document.paragraphs[0].text
    current_entry = ''
    for index, paragraph in enumerate(document.paragraphs[1:]):
        if date_match(paragraph.text) and index > 0:
            if(current_entry != ''):
                cursor.execute("INSERT INTO entries VALUES (?, ?)", (current_date, current_entry.decode('utf-8')))
            current_date = paragraph.text
            current_entry = ''
        else:
            current_entry += '\n' + paragraph.text.encode('utf-8')

# insert the last entry into the database
cursor.execute("INSERT INTO entries VALUES ('{}', '{}')".format(current_date, current_entry))

cursor.execute("SELECT COUNT(*) FROM entries")
# print(cursor.fetchall())

conn.commit()
conn.close()
